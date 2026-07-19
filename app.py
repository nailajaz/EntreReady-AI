from docx import Document
from docx.shared import Pt

doc=Document()
doc.add_heading('EntreReady AI (ERAI)\nResearch Design Notes',1)
p=doc.add_paragraph()
p.add_run("Key Principle\n").bold=True
p.add_run("Because the Qualtrics data has already been collected, the original survey instrument must remain the empirical foundation of the research. The prototype should not replace or contradict the collected questionnaire.\n")

doc.add_heading("Two-Layer Framework", level=2)

doc.add_heading("Layer 1 – Research Validation (Existing Dataset)", level=3)
items=[
"Use the existing Qualtrics dataset (214 responses) as the research dataset.",
"Train and validate the machine learning models using the collected data.",
"Evaluate prediction performance and compare algorithms.",
"Identify significant predictors and explain model outputs.",
"Present this as the empirical methodology in the journal paper."
]
for i in items:
    doc.add_paragraph(i, style="List Bullet")

doc.add_heading("Layer 2 – Live Prototype (EntreReady AI – ERAI)", level=3)
doc.add_paragraph(
"The Streamlit prototype is an operational implementation of the validated framework. "
"It should not claim to reproduce the original research questionnaire exactly."
)

doc.add_paragraph(
'Suggested wording for the paper:', style=None).bold=True
doc.add_paragraph(
'"The assessment interface is an operational implementation of the EntreReady AI framework. '
'It translates the validated constructs into a practical decision-support tool for entrepreneurship education and entrepreneurial development."'
)

doc.add_heading("Prototype Design Approach", level=2)
doc.add_paragraph(
"Instead of presenting all research survey questions, the prototype should use a shorter, user-friendly assessment "
"(approximately 15–20 questions) that maps to the validated constructs."
)

table=doc.add_table(rows=1, cols=2)
table.style='Table Grid'
hdr=table.rows[0].cells
hdr[0].text='Research Construct'
hdr[1].text='Example Prototype Question'
rows=[
("Entrepreneurial Self-Efficacy","How confident are you in leading a business?"),
("AI Usage","How often do you use AI tools to solve business problems?"),
("Entrepreneurial Intention","Do you plan to start a business within the next five years?")
]
for a,b in rows:
    c=table.add_row().cells
    c[0].text=a
    c[1].text=b

doc.add_heading("How to Present This in the Paper", level=2)
for t in [
"Research Phase: Survey instrument, statistical validation, and machine learning training using the collected Qualtrics dataset.",
"Implementation Phase: EntreReady AI prototype, simplified assessment interface, Explainable AI, and personalized recommendations."
]:
    doc.add_paragraph(t, style="List Bullet")

doc.add_heading("Recommendation", level=2)
doc.add_paragraph(
"Before finalizing the prototype questionnaire, map each existing Qualtrics construct to one or more practical assessment questions. "
"This ensures consistency between the empirical research and the operational AI prototype while maintaining academic rigor."
)

out="/mnt/data/EntreReady_AI_ERAI_Research_Design_Notes.docx"
doc.save(out)
print(out)
